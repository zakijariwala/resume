#!/usr/bin/env python3
"""
Generate the AI Product Manager resume as a clean, ATS-friendly .docx.
PDF is produced from the .docx via LibreOffice (soffice --convert-to pdf).

This is a standalone generator seeded from the portfolio content. It is the
starting point for the automated PDF generator to be wired into the site later.

Run:  python3 misc/generate_ai_pm_resume.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Theme ──────────────────────────────────────────────────────────────────
ACCENT = RGBColor(0x1F, 0x2A, 0x44)   # dark navy headings
MUTED  = RGBColor(0x4B, 0x55, 0x63)   # muted grey
BODY   = RGBColor(0x1A, 0x1A, 0x1A)
FONT   = "Calibri"
PAGE_W = 7.5  # usable width in inches with 0.5" margins on Letter

# ── Content ────────────────────────────────────────────────────────────────
NAME    = "Mohammad Zaki Jariwala"
TITLE   = "AI Product Manager · Technical"
CONTACT = [
    "Navi Mumbai, India · Open to Global Relocation",
    "+91 9601406034  |  Jariwalazaki@gmail.com",
    "linkedin.com/in/zakijariwala  |  github.com/zakijariwala",
]

SUMMARY = (
    "Technical Product Manager specialising in AI-native products and automation workflows. "
    "Three years leading cross-functional delivery at State Bank of India — one of India's "
    "largest banks — managing platforms serving 30,000+ users while driving operational "
    "modernisation. Proven 0-to-1 product ownership across consumer apps, B2B tools, and "
    "LLM-powered pipelines — from user research and discovery through to production. "
    "Technical depth in LLM orchestration, Python, and API integration enables direct "
    "collaboration with engineering teams without translation overhead."
)

SKILLS = [
    ("Product",
     "Product Discovery & User Research · Roadmapping & Prioritisation (RICE/MoSCoW) · "
     "AI Product Strategy · Product Analytics (PostHog, funnel analysis) · Go-to-Market "
     "Execution · Stakeholder & Executive Communication · PRD Writing · Cross-functional Leadership"),
    ("AI & LLM",
     "Anthropic Claude API · Google Gemini · Prompt Engineering · LLM Orchestration & "
     "Governance · GenAI Workflow Design · AI Agent Architecture"),
    ("Technical",
     "Python (ETL, pipelines, scripting) · JavaScript/TypeScript · SQLite · "
     "GitHub Actions & CI/CD · Docker · Cloudflare"),
    ("Infrastructure",
     "Linux (RHEL/Ubuntu) · Identity & Access Management · High Availability Systems "
     "(99.999% SLA) · Disaster Recovery & Incident Management"),
]

EXPERIENCE = [
    {
        "org": "Tata Consultancy Services · State Bank of India (SBI GITC)",
        "date": "2023 — Present",
        "role": "Systems Engineer, IDM & Infrastructure",
        "metrics": "99.999% Uptime · 30,000+ Users · 150+ Servers · 50% RTO Reduction",
        "bullets": [
            "Led an operations team delivering 99.999% availability for Tier-1 banking applications across 30,000+ branch users — setting team priorities, mentoring 5 engineers, and managing escalations end-to-end.",
            "Collaborated across infrastructure, application, database, and security teams to align delivery priorities and maintain regulatory compliance under strict banking constraints.",
            "Drove infrastructure modernisation from concept to delivery — evaluated build-vs-buy tradeoffs, secured stakeholder alignment, and led adoption of on-premise Docker workflows within compliance boundaries.",
            "Architected and executed Disaster Recovery runbooks in cross-functional coordination with application and database teams, reducing Recovery Time Objective (RTO) by 50%.",
            "Translate complex incident data and technical risk into clear communications for business stakeholders and senior leadership.",
        ],
    },
    {
        "org": "Independent Coaching Practice",
        "date": "2025 — 2026 (Concluded)",
        "role": "Founder & Technical Communication Coach",
        "metrics": "50+ Clients · 1:1 Format",
        "bullets": [
            "Founded and scaled a coaching practice from zero — handling client acquisition, curriculum design, and end-to-end delivery for 50+ engineering and management professionals.",
            "Designed a structured curriculum covering STAR-method framing, stakeholder communication, and senior-level interview strategy — iterating based on client outcomes and session feedback.",
            "Developed repeatable frameworks for technical storytelling, enabling clients to translate complex engineering work into language accessible to non-technical decision-makers.",
        ],
    },
    {
        "org": "Zamaan Marine",
        "date": "2025 — Present",
        "role": "Technology & Product Lead",
        "metrics": "Family marine-equipment venture · Unpaid  |  200+ B2B Prospects/Campaign · 90%+ Inbox Delivery · ~60% Faster Listing",
        "bullets": [
            "Identified the digital gap in a family marine-parts business through direct conversations with 200+ B2B prospects — validating channels, value proposition, and pricing signals before building.",
            "Defined and prioritised a zero-budget digital stack: JAMstack storefront for discoverability, eBay API automation for listings (~60% faster time-to-market), and an AI-powered cold-email funnel for B2B acquisition.",
            "Instrumented product analytics via PostHog, tracking the full funnel from storefront visit to eBay conversion — using data to reprioritise listing categories and outreach targeting.",
            "Shipped an AI cold-email agent (Claude API + Python + Brevo) achieving 90%+ inbox delivery across campaign cycles.",
        ],
    },
]

PROJECTS = [
    {
        "name": "Content Automation Pipeline",
        "date": "2025 — Present",
        "tag": "End-to-End LLM Orchestration · 85% Faster Drafts · 12 Niches",
        "stack": "Python · Gemini 1.5 Pro · Tauri · React 19 · GitHub Actions",
        "bullets": [
            "Identified that content creators spend most research-to-draft time on mechanical, repeatable work — designed and shipped a pipeline automating transcript ingestion, LLM transformation, and publishing.",
            "Defined and built an LLM governance layer (\"Failure First\" framework) to solve a specific output-quality problem — eliminating safe, generic AI prose across 12 distinct content niches.",
            "Reduced time-to-draft by 85%; built a cross-platform desktop GUI (Tauri/React 19) enabling non-technical users to operate the pipeline without CLI access.",
        ],
    },
    {
        "name": "Path of Supplication — Islamic Supplications PWA",
        "date": "2024 — Present",
        "tag": "Offline-First PWA · 40+ Beta Testers · Sub-10ms Queries",
        "stack": "Python · JavaScript · PWA · Service Workers · SQLite",
        "bullets": [
            "Identified a gap in available supplication apps — existing options were ad-heavy, required constant connectivity, or lacked multilingual support — and validated the problem before building.",
            "Structured a publicly available liturgical database of 22,000+ records into an optimised, indexed SQLite schema, achieving sub-10ms local query times — 4× faster than existing alternatives.",
            "Shipped as an offline-first PWA (Service Workers, lazy caching), delivering 80MB of multilingual content (Arabic, Urdu, English) at zero hosting cost.",
            "Ran a structured beta with 40+ testers across iOS Safari, Android Chrome, and desktop — triaging feedback and shipping iterative improvements based on user-reported priorities.",
        ],
    },
    {
        "name": "run.to — Activity Platform",
        "date": "2025",
        "tag": "Consumer Mobile App · 50+ Alpha Users · Flutter",
        "stack": "Flutter · Dart · Supabase · PostGIS · Riverpod",
        "bullets": [
            "Built a consumer running app (GPS tracking, ghost routes, crowd radar) end-to-end across 6 feature segments — owning architecture, backend, and UX decisions.",
            "Grew to 50+ local alpha users in early testing, validating core flows and GPS accuracy before deliberately taking the product private to polish based on feedback.",
            "Integrated Supabase + PostGIS for geospatial features, supporting sub-100ms proximity queries for nearby-runner detection.",
        ],
    },
]

EDUCATION = {
    "degree": "B.Tech — Electronics & Communication Engineering",
    "institution": "Dharmsinh Desai University",
    "year": "Graduated 2023",
}

CERTS = [
    "AWS Certified Solutions Architect – Associate (SAA-C03) — Amazon Web Services · 2024",
    "Google GenAI Leader — Google Cloud · 2025",
]


# ── Helpers ────────────────────────────────────────────────────────────────
def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)
    style.font.color.rgb = BODY
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.line_spacing = 1.05


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT
    r.font.name = FONT
    # bottom border
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "1F2A44")
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)
    return p


def add_entry_header(doc, left, right):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(PAGE_W), WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run(left)
    r.bold = True
    r.font.size = Pt(10.5)
    r2 = p.add_run("\t" + right)
    r2.font.color.rgb = MUTED
    r2.font.size = Pt(10)
    return p


def add_subline(doc, role, metrics=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(role)
    r.italic = True
    r.font.color.rgb = MUTED
    r.font.size = Pt(10)
    if metrics:
        r2 = p.add_run("   " + metrics)
        r2.font.color.rgb = MUTED
        r2.font.size = Pt(9)
    return p


def add_bullets(doc, bullets):
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.line_spacing = 1.04
        r = p.add_run(b)
        r.font.size = Pt(10)


# ── Build ──────────────────────────────────────────────────────────────────
def build(out_docx):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.5)
        s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.5)
        s.right_margin = Inches(0.5)
    set_base_style(doc)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(NAME)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(TITLE)
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED

    for i, line in enumerate(CONTACT):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1 if i < len(CONTACT) - 1 else 4)
        r = p.add_run(line)
        r.font.size = Pt(9.5)
        r.font.color.rgb = MUTED

    # Summary
    add_section_heading(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(SUMMARY)
    r.font.size = Pt(10)

    # Skills
    add_section_heading(doc, "Core Competencies")
    for cat, items in SKILLS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        rc = p.add_run(cat + ":  ")
        rc.bold = True
        rc.font.size = Pt(10)
        ri = p.add_run(items)
        ri.font.size = Pt(10)

    # Experience
    add_section_heading(doc, "Professional Experience")
    for e in EXPERIENCE:
        add_entry_header(doc, e["org"], e["date"])
        add_subline(doc, e["role"], e["metrics"])
        add_bullets(doc, e["bullets"])

    # Projects
    add_section_heading(doc, "Selected Projects")
    for pr in PROJECTS:
        add_entry_header(doc, pr["name"], pr["date"])
        add_subline(doc, pr["tag"], None)
        ps = doc.add_paragraph()
        ps.paragraph_format.space_after = Pt(2)
        rs = ps.add_run("Stack: " + pr["stack"])
        rs.font.size = Pt(9)
        rs.font.color.rgb = MUTED
        add_bullets(doc, pr["bullets"])

    # Education & Certs
    add_section_heading(doc, "Education & Certifications")
    p = add_entry_header(doc, EDUCATION["degree"], EDUCATION["year"])
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(3)
    r2 = p2.add_run(EDUCATION["institution"])
    r2.italic = True
    r2.font.color.rgb = MUTED
    r2.font.size = Pt(10)
    add_bullets(doc, CERTS)

    doc.save(out_docx)
    print("Wrote", out_docx)


if __name__ == "__main__":
    import sys
    out = sys.argv[1] if len(sys.argv) > 1 else "Mohammad_Zaki_Jariwala_Resume.docx"
    build(out)
