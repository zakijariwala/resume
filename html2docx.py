import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    
    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0000FF')
    rPr.append(c)

    # Add underline
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def create_docx_from_html(html_path, docx_path):
    import docx
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')

    # HEADER
    name_el = soup.find(class_='name')
    if name_el:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name_el.get_text(strip=True))
        run.bold = True
        run.font.size = Pt(24)

    title_el = soup.find(class_='title-tag')
    if title_el:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_el.get_text(strip=True))
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x4b, 0x55, 0x63)

    contact_el = soup.find(class_='contact')
    if contact_el:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        contact_lines = [line.strip() for line in contact_el.strings if line.strip() and line.strip() != '·']
        run = p.add_run(" | ".join(contact_lines))
        run.font.size = Pt(10)

    # SECTIONS
    for section in soup.find_all('section'):
        # Add some space before section
        doc.add_paragraph()
        
        label_el = section.find(class_='section-label')
        if label_el:
            p = doc.add_heading(label_el.get_text(strip=True).upper(), level=1)
            # Add bottom border to heading (requires oxml)
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '000000')
            pBdr.append(bottom)
            p._p.get_or_add_pPr().append(pBdr)

        # SUMMARY
        for summary in section.find_all(class_='summary', recursive=False):
            p = doc.add_paragraph(summary.get_text(strip=True))
            p.paragraph_format.space_after = Pt(12)

        # ENTRIES
        for entry in section.find_all(class_='entry'):
            header_el = entry.find(class_='entry-header')
            org_el = entry.find(class_='entry-org') or entry.find(class_='entry-subtitle')
            bullets_el = entry.find(class_='bullets')
            summary_el = entry.find(class_='summary')

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            
            if header_el:
                title = header_el.find(class_='entry-title')
                date = header_el.find(class_='entry-date')
                
                if title:
                    r = p.add_run(title.get_text(strip=True))
                    r.bold = True
                
                if date:
                    # add right aligned tab for date
                    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
                    r = p.add_run("\t" + date.get_text(strip=True))
            
            if org_el:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(4)
                r2 = p2.add_run(org_el.get_text(strip=True))
                r2.italic = True
                
            if bullets_el:
                for li in bullets_el.find_all('li'):
                    p_li = doc.add_paragraph(li.get_text(strip=True), style='List Bullet')
                    p_li.paragraph_format.space_after = Pt(2)
                    
            if summary_el:
                p_sum = doc.add_paragraph(summary_el.get_text(strip=True))
                p_sum.paragraph_format.space_after = Pt(6)
        
        # SKILLS (handle specifically)
        skills_grid = section.find(class_='skills-grid') or section.find(class_='skills-container') or section.find(class_='skills-text')
        if skills_grid:
            if 'skills-text' in skills_grid.get('class', []):
                # Paragraph based skills
                for span in skills_grid.find_all('span', class_='skills-category'):
                    p = doc.add_paragraph()
                    r = p.add_run(span.get_text(strip=True) + " ")
                    r.bold = True
                    # Get the following text sibling
                    next_node = span.next_sibling
                    if next_node and isinstance(next_node, str):
                        p.add_run(next_node.strip())
            else:
                for row in skills_grid.find_all(class_='skill-row'):
                    cat = row.find(class_='skill-cat') or row.find(class_='skill-category')
                    items = row.find(class_='skill-items')
                    
                    if cat:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(2)
                        r = p.add_run(cat.get_text(strip=True) + ": ")
                        r.bold = True
                        
                        text_part = ""
                        if items:
                            text_part = items.get_text(strip=True)
                        else:
                            # if items doesn't exist, the text might be directly in the row
                            text_part = row.get_text(strip=True).replace(cat.get_text(strip=True), "").strip()
                        
                        p.add_run(text_part)

        # CERTS (handle specifically)
        certs_list = section.find(class_='certs-list')
        if certs_list:
            for item in certs_list.find_all(['div'], class_=['cert-item', 'cert-entry']):
                name = item.find(class_='cert-name')
                org = item.find(class_='cert-org')
                
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                if name:
                    r = p.add_run(name.get_text(strip=True) + " - ")
                    r.bold = True
                if org:
                    p.add_run(org.get_text(strip=True))

    doc.save(docx_path)
    print(f"Created {docx_path}")

files = {
    'v_sre.html': 'Mohammad_Zaki_Jariwala_Resume_SRE.docx',
    'v_tpm.html': 'Mohammad_Zaki_Jariwala_Resume_TPM.docx',
    'v_mba.html': 'Mohammad_Zaki_Jariwala_Resume_MBA.docx',
    'v_ai.html': 'Mohammad_Zaki_Jariwala_Resume_AI.docx'
}

base_dir = r"D:\Zaki\Projects\resume"
for html_file, docx_name in files.items():
    html_path = os.path.join(base_dir, html_file)
    docx_path = os.path.join(base_dir, docx_name)
    if os.path.exists(html_path):
        create_docx_from_html(html_path, docx_path)
